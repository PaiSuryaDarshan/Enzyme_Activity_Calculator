# Bradford Assay Functions

from docx import Document
import Pi_Bradford as PB
from datetime import datetime

def EditTable(doc, df, index):
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    print(f"Table at Position {index}")

    t = doc.tables[index]

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

class Brad:

    def Edit_Brad_temp(df_norm, df_calc, trend, ImageName, savename = f"Brad_{datetime.now().date()}.docx"):

        doc = Document("./Templates/Bradfor_Temp.docx")

        dict_to_replace = {
            '{Date}': f'Date: {datetime.now().date()}',
            '{Time}': f'Time: {datetime.now().time()}',
            '{Image1}': "PLACEHOLDER",
            '{Image2}': "PLACEHOLDER",
            '{Graph_Eq}': f"y = {trend.slope.round(4)}x + {trend.intercept.round(4)} \n R2 = {trend.rvalue.round(4)}",
            }            
        
        counter = 0
        for paragraph in doc.paragraphs:
            for key in dict_to_replace.keys():
                if f"{key}" in paragraph.text and dict_to_replace[key] != "PLACEHOLDER":
                    print(paragraph.text)
                    paragraph.text = dict_to_replace[key]
                    break

                elif f"{key}" in paragraph.text and f"{key}" == "{Image1}":
                    print (paragraph.text)
                    p = paragraph
                    p.text = "Bradford Assay Curve"
                    r = p.add_run()
                    imagePath = "./Data_Storage/"
                    r.add_picture(imagePath+ImageName+' w.o. Unknown.png')
                    break

                elif f"{key}" in paragraph.text and f"{key}" == "{Image2}":
                    print (paragraph.text)
                    paragraph.text = "Bradford Assay with Unknown(s)"
                    p = paragraph
                    r = p.add_run()
                    imagePath = "./Data_Storage/"
                    r.add_picture(imagePath+ImageName+' w Unknown.png')
                    break
        
        EditTable(doc, df_norm, 0)               
        EditTable(doc, df_calc, 1)               

        doc.save(f"./Output/{savename}")
        
        return

class EnzAssay:

    def Edit_Enz_temp(ImageName, savename = f"Brad_{datetime.now().date()}.docx"):

    #     doc = Document("./Templates/Bradfor_Temp.docx")

    #     dict_to_replace = {
    #         '{Date}': f'Date: {datetime.now().date()}',
    #         '{Time}': f'Time: {datetime.now().time()}',
    #         '{Image1}': "PLACEHOLDER",
    #         '{Image2}': "PLACEHOLDER",
    #         '{Graph_Eq}': f"y = {trend.slope.round(4)}x + {trend.intercept.round(4)} \n R2 = {trend.rvalue.round(4)}",
    #         }            
        
    #     counter = 0
    #     for paragraph in doc.paragraphs:
    #         for key in dict_to_replace.keys():
    #             if f"{key}" in paragraph.text and dict_to_replace[key] != "PLACEHOLDER":
    #                 print(paragraph.text)
    #                 paragraph.text = dict_to_replace[key]
    #                 break

    #             elif f"{key}" in paragraph.text and f"{key}" == "{Image1}":
    #                 print (paragraph.text)
    #                 p = paragraph
    #                 p.text = "Bradford Assay Curve"
    #                 r = p.add_run()
    #                 imagePath = "./Data_Storage/"
    #                 r.add_picture(imagePath+ImageName+' w.o. Unknown.png')
    #                 break

    #             elif f"{key}" in paragraph.text and f"{key}" == "{Image2}":
    #                 print (paragraph.text)
    #                 paragraph.text = "Bradford Assay with Unknown(s)"
    #                 p = paragraph
    #                 r = p.add_run()
    #                 imagePath = "./Data_Storage/"
    #                 r.add_picture(imagePath+ImageName+' w Unknown.png')
    #                 break
        
    #     EditTable(doc, df_norm, 0)               
    #     EditTable(doc, df_calc, 1)               

    #     doc.save(f"./Output/{savename}")
        
        return